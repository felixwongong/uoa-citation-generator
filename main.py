from newspaper import Article
from docx import Document

from datetime import datetime
import sys

citeDictArr = []
citeDict = {}

isReplace = False


document = Document(sys.argv[1] if len(sys.argv) > 1 else None)

while True:
    inputUrl = input(
        "Please input your article url. If this is the last one, type quit\n")
    if(inputUrl.strip() == 'quit'):
        break
    else:
        article = Article(inputUrl)

        article.download()
        article.parse()
        citeDict['author'] = ", ".join(article.authors)
        citeDict['title'] = article.title
        citeDict['date'] = article.publish_date
        citeDict['url'] = article.url
        citeDict['dateAccess'] = datetime.today().strftime('%Y, %B %d')
        citeDictArr.append(citeDict)


for i in range(len(citeDictArr)):
    dict = citeDictArr[i]
    author = f'{dict["author"]}. ' if dict['author'] != '' else ''
    date = f'({dict["date"]})' if dict['date'] else ''
    p = document.add_paragraph(f'{author}', style="List Number")
    p.add_run(f'{citeDictArr[i]["title"]} ').italic = True
    p.add_run(
        f'{date} [Online]. Available: {dict["url"]} [{dict["dateAccess"]}]\n')


document.save(sys.argv[1] if len(sys.argv) > 1 else "reference.docx")
